from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PATCH_VERSION = "v0.6.1"
DOC_PATHS = [
    ROOT / "物价补丁" / "使用文档.docx",
    ROOT / "发布版" / "物价补丁" / "使用文档.docx",
]
OUT_DIR = ROOT / "output" / "doc"
TEMPLATE_DOC = ROOT / "docs" / "使用文档.docx"


def copy_template_doc() -> bool:
    if not TEMPLATE_DOC.exists():
        return False

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE_DOC)
    doc.core_properties.title = f"POE1 / POE2 物价补丁使用文档 {PATCH_VERSION}"
    replacements = {
        "POE2 三服合一物价补丁使用文档": (
            f"POE1 / POE2 物价补丁使用文档 {PATCH_VERSION}"
        ),
        "程序会自动识别国服 WeGame、国际服官方 GGPK、国际服 Steam/Epic Bundles2。": (
            "更新和还原都会自动识别 POE1 / POE2、国服 WeGame、国际服官方 GGPK、国际服 Steam/Epic Bundles2，也可以手动选择游戏目录。"
        ),
        "更新和还原都会自动识别国服 WeGame、国际服官方 GGPK、国际服 Steam/Epic Bundles2，也可以手动选择游戏目录。": (
            "更新和还原都会自动识别 POE1 / POE2、国服 WeGame、国际服官方 GGPK、国际服 Steam/Epic Bundles2，也可以手动选择游戏目录。"
        ),
        "把整个“物价补丁”文件夹放到 POE2 游戏根目录。脚本会自动识别上一级目录，不写死盘符。": (
            "“物价补丁”文件夹可以放在任意位置。放在游戏根目录下仍可最快识别；放在其它位置时，程序会检查最近目录、GGG 官服注册表和各平台游戏库，也可以手动浏览。"
        ),
        "“物价补丁”文件夹可以放在任意位置。放在游戏根目录下仍可最快识别；放在其它位置时，程序会检查最近目录和各平台游戏库，也可以手动浏览。": (
            "“物价补丁”文件夹可以放在任意位置。放在游戏根目录下仍可最快识别；放在其它位置时，程序会检查最近目录、GGG 官服注册表和各平台游戏库，也可以手动浏览。"
        ),
        "二、自动识别": "二、选择与自动识别游戏目录",
        r"<POE2游戏根目录>\Content.ggpk": r"<POE1 或 POE2 游戏根目录>\Content.ggpk",
        r"<POE2游戏根目录>\Bundles2\_.index.bin": r"<POE1 或 POE2 游戏根目录>\Bundles2\_.index.bin",
        "检测到 Content.ggpk：按国际服官方 GGPK 处理。": (
            "自动模式依次检查补丁文件夹上一级、环境变量、最近有效目录、GGG 官服注册表、已安装程序、WeGame/Steam 游戏库、Epic 清单和常见位置。"
        ),
        "自动模式依次检查补丁文件夹上一级、环境变量、最近有效目录、已安装程序、WeGame/Steam 游戏库、Epic 清单和常见位置。": (
            "自动模式依次检查补丁文件夹上一级、环境变量、最近有效目录、GGG 官服注册表、已安装程序、WeGame/Steam 游戏库、Epic 清单和常见位置。"
        ),
        "检测到 Bundles2 且有 WeGame/腾讯文件特征：按国服 WeGame Bundles2 处理。": (
            "国服会匹配“流放之路：降临”和 WeGameApps\\rail_apps 游戏库；目录还必须包含 Bundles2\\_.index.bin，并通过 WeGame/Rail/Tencent 特征确认。"
        ),
        "检测到 Bundles2 且没有 WeGame/腾讯文件特征：按国际服 Steam/Epic Bundles2 处理。": (
            "检测到 Content.ggpk 时按国际服官方 GGPK 处理；检测到非国服特征的 Bundles2 时按国际服 Steam/Epic 处理。"
        ),
        "国际服会读取当前游戏 language 设置，自动写入对应语言的 BaseItemTypes。": (
            "如果发现多个客户端，程序不会猜测，请切换到手动选择。确认后的有效目录会保存，更新与还原下次可复用。POE1 还可在 GUI 选择自动识别、汉化补丁、简体中文、繁体中文或跟随游戏配置；自动模式会用最新客户端日志识别第三方汉化补丁。"
        ),
        "如果发现多个客户端，程序不会猜测，请切换到手动选择。确认后的有效目录会保存，更新与还原下次可复用。国际服仍会读取 language 设置。": (
            "如果发现多个客户端，程序不会猜测，请切换到手动选择。确认后的有效目录会保存，更新与还原下次可复用。POE1 还可在 GUI 选择自动识别、汉化补丁、简体中文、繁体中文或跟随游戏配置；自动模式会用最新客户端日志识别第三方汉化补丁。"
        ),
        "3. 程序会提取英文表和当前客户端目标语言表。": (
            "3. 在窗口中确认自动识别的路径和客户端类型，或切换到手动选择；POE1 还需确认显示语言。价格赛季位于窗口顶部“游戏版本”右侧，自动识别时会随客户端下拉框切换 POE1/POE2 赛季（默认最新，可点击刷新发现新赛季）。POE1/POE2 国服只支持当前赛季。更新与还原必须使用同一目标语言。"
        ),
        "3. 在窗口中确认自动识别的路径和客户端类型，或切换到手动选择；更新器还需选择补丁范围。随后程序提取英文表和当前客户端目标语言表。": (
            "3. 在窗口中确认自动识别的路径和客户端类型，或切换到手动选择；POE1 还需确认显示语言。价格赛季位于窗口顶部“游戏版本”右侧，自动识别时会随客户端下拉框切换 POE1/POE2 赛季（默认最新，可点击刷新发现新赛季）。POE1/POE2 国服只支持当前赛季。更新与还原必须使用同一目标语言。"
        ),
        "2. 双击“一键还原物价补丁.exe”。": (
            "2. 双击“物价补丁.exe”，确认自动识别的路径和客户端类型，或手动选择要还原的游戏目录；点击底部“还原物价补丁”执行还原。"
        ),
        "2. 双击“一键更新物价补丁.exe”。": (
            "2. 双击“物价补丁.exe”，确认自动识别的路径和客户端类型，或手动选择要更新的游戏目录；点击底部“开始/更新物价补丁”执行更新。"
        ),
        "2. 双击“一键还原物价补丁.exe”，在还原窗口中确认自动识别的路径和客户端类型，或手动选择要还原的游戏目录。": (
            "2. 双击“物价补丁.exe”，确认自动识别的路径和客户端类型，或手动选择要还原的游戏目录；点击底部“还原物价补丁”执行还原。"
        ),
        "三、一键启动/更新": "三、统一操作入口",
        "四、一键还原": "四、还原物价补丁",
        r"<POE2游戏根目录>\物价补丁\一键更新物价补丁.exe": (
            r"<POE 游戏根目录>\物价补丁\物价补丁.exe"
        ),
        r"<POE2游戏根目录>\物价补丁\一键还原物价补丁.exe": (
            r"<POE 游戏根目录>\物价补丁\物价补丁.exe"
        ),
        "一键更新物价补丁.exe：抓价、生成补丁并写入游戏包。": (
            "物价补丁.exe：统一 GUI 入口；可选择更新或还原、POE1 或 POE2、自动识别或手动路径。"
        ),
        "一键还原物价补丁.exe：还原对应 BaseItemTypes。": (
            "物价补丁.exe：统一 GUI 入口；可选择更新或还原、POE1 或 POE2、自动识别或手动路径。"
        ),
        "4. 程序会抓取 poe2scout 国际服价格，并把价格追加为“=数字D/E”。": (
            "4. 程序会按版本、客户端类型和所选价格赛季抓取价格：POE1 国际服使用 poe.ninja 主源，再由 poe2scout 和 PoEDB 逐物品补缺；"
            "POE1 国服使用 poecurrency.top 主源，再按 poe.ninja、poe2scout、PoEDB 的顺序补缺；POE2 数据源和 D/E 规则保持不变。"
            "各来源独立失败和降级，POE1 使用 C/D，POE2 使用 D/E，比例使用当前数据源实时值。"
            "同一个显示名对应多个游戏元数据路径时会为全部别名写入价格；国服翻译重名时使用 engname 英文别名消歧，避免漏价或串价。"
        ),
        "5. 程序会生成“物价补丁.zip”和“还原物价补丁.zip”。Bundles2 模式还会生成“真实还原物价补丁.zip”。": (
            "5. 程序会生成物价补丁和按 POE1/POE2、客户端类型、目标语言独立命名的还原基线；工作输出再按游戏目录哈希隔离。"
            "历史赛季使用独立的赛季标识和缓存；权威副本只保存在对应游戏根目录的 .poe1-price-patch 或 .poe2-price-patch，官方 GGPK、Steam/Epic 与国服不会互相覆盖；连续更新只复用第一次验证通过的干净基线。"
        ),
        "6. 程序会把补丁写回对应游戏包。": (
            "6. 价格文件会先在隔离目录生成并校验；实时数据失败时仅使用当前范围和模式的兼容核心缓存。"
            "写入前复核游戏文件未被并发修改，写入后读回核对；失败会自动恢复。没有安全候选时只停止本次更新，不覆盖当前游戏。"
        ),
        "3. Bundles2 模式会使用“真实还原物价补丁.zip”恢复打补丁前的物理文件。": (
            "3. Bundles2 模式优先验证并使用“真实还原物价补丁.zip”，无需先准备 .NET；替换后核对精确文件集合，失败会恢复操作前状态。"
        ),
        "4. GGPK 模式会使用“还原物价补丁.zip”写回当前客户端对应的 BaseItemTypes。": (
            "4. GGPK 模式会使用“还原物价补丁.zip”写回当前语言目标，随后逐条读回校验；写入或校验失败会自动重试一次。"
        ),
        "5. 如果没有可用还原包，程序会拒绝做不完整还原。": (
            "5. 如果专属还原基线丢失或过期，程序会在临时目录只清除能够确认属于本工具的 BaseItemTypes、Words 与岛屿提示，结构和零残留校验通过后自动重建；只有无法证明安全时才会拒绝并提示平台修复。"
        ),
        "Bundles2 模式不覆盖完整 _.index.bin 或 Tiny*.bundle.bin。": (
            "Bundles2 模式会更新 _.index.bin 并向 LibGGPK3 写入增量内容，不会直接覆盖 Tiny*.bundle.bin。"
        ),
        "Bundles2 还原会恢复安装前备份的 _.index.bin 和 LibGGPK3 状态。": (
            "Bundles2 还原优先恢复安装前备份的 _.index.bin 和 LibGGPK3；迁移基线则恢复到语义上干净的物价层。"
        ),
        "真实还原物价补丁.zip：Bundles2 模式的物理级恢复包。": (
            "POE1/POE2 专属还原包：名称包含客户端类型和语言；Bundles2 还会保存物理恢复包。持久副本分别位于 <游戏根目录>\\.poe1-price-patch 与 <游戏根目录>\\.poe2-price-patch。"
        ),
        "国服 WeGame：data/balance/simplified chinese/baseitemtypes.datc64。": (
            "国服 WeGame：POE1 写入 data/simplified chinese/baseitemtypes.datc64，POE2 写入 data/balance/simplified chinese/baseitemtypes.datc64。"
        ),
        "国际服官方 / Steam / Epic：按当前游戏语言写入，例如繁中为 data/balance/traditional chinese/baseitemtypes.datc64，英文为 data/balance/baseitemtypes.datc64。": (
            "国际服官方 / Steam / Epic：POE1 繁中路径为 data/traditional chinese/baseitemtypes.datc64；POE2 繁中路径为 data/balance/traditional chinese/baseitemtypes.datc64。其它语言按各自实际资源表写入。"
        ),
        "需要手动指定语言时，可设置 POE2_PATCH_LANGUAGE，例如 zh-TW、en、ja。": (
            "POE1 在 GUI 选择自动识别、汉化补丁、简中、繁中或跟随配置；POE2 可设置 POE2_PATCH_LANGUAGE，例如 zh-TW、en、ja。"
        ),
        r"tools\dotnet-runtime：内置 .NET 8 runtime，不要删除。": (
            r"tools\dotnet-runtime：内置 .NET 8.0.28 runtime，不要删除；发布包另含校验过的离线修复包，正常使用不需要联网，只有两者都不可用时才访问 Microsoft 源。"
        ),
        r"tools\python：内置 Python 和依赖，不要删除。": (
            r"tools\python：内置 Python 和依赖，不要删除；自动修复时依次尝试华为云、阿里云、南京大学镜像和 Python 官方备用源。"
        ),
        "提示找不到游戏目录：请确认物价补丁文件夹放在 POE2 游戏根目录。": (
            "提示找不到游戏目录：改用手动选择，并选择直接包含 Content.ggpk 或 Bundles2\\_.index.bin 的游戏根目录；不要选择物价补丁文件夹。多个客户端必须明确选择。若旧补丁有标记但备份丢失，新版会搜索旧文件夹并尝试离线沙盒迁移；失败时真实游戏不会被修改。"
        ),
        "提取或写入失败：请先关闭游戏和可能占用文件的工具。": (
            "提取或写入失败：程序会等待短暂占用并在写入失败时自动恢复；仍失败时请关闭游戏和可能占用文件的工具后重试。"
        ),
        "缺少价格：可能是 poe2scout 暂无该物品数据，或英文名无法匹配本地物品表。": (
            "缺少价格：可能是当前数据源暂无该物品，或名称无法匹配本地表。POE1 安装第三方汉化补丁时，请把显示语言设为自动识别或汉化补丁；低匹配会保留未命中的旧价格，数据源完全不可用时保留当前补丁。POE1 传奇写成 传奇名[<<价格>>]，POE2 写成 [价格|传奇名]；易刷与 PoE Overlay II 会按对应格式清理价格后使用精确传奇名查询。"
        ),
    }
    matched: set[str] = set()
    title_key = "POE2 三服合一物价补丁使用文档"
    for paragraph in doc.paragraphs:
        original = paragraph.text.strip()
        if original == title_key or original.startswith(f"{title_key} v"):
            replacement = f"POE1 / POE2 物价补丁使用文档 {PATCH_VERSION}"
            matched.add(title_key)
        else:
            replacement = replacements.get(original)
        if replacement is None:
            continue
        if original != title_key and not original.startswith(f"{title_key} v"):
            matched.add(original)
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replacement)
    current_text = {paragraph.text.strip() for paragraph in doc.paragraphs}
    missing = sorted(
        original
        for original, replacement in replacements.items()
        if original not in matched and replacement not in current_text
    )
    if missing:
        raise RuntimeError(
            "release document template no longer contains expected paragraphs: "
            + " | ".join(missing)
        )

    duplicate_texts = {
        r"<POE 游戏根目录>\物价补丁\物价补丁.exe",
        "物价补丁.exe：统一 GUI 入口；可选择更新或还原、POE1 或 POE2、自动识别或手动路径。",
    }
    seen_duplicates: set[str] = set()
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text not in duplicate_texts:
            continue
        if text in seen_duplicates:
            element = paragraph._element
            element.getparent().remove(element)
        else:
            seen_duplicates.add(text)

    # The WPS-authored template's contextualSpacing flag makes some
    # LibreOffice versions overlap consecutive list paragraphs after wrapping.
    list_style_ppr = doc.styles["List Bullet"]._element.get_or_add_pPr()
    for node in list_style_ppr.findall(qn("w:contextualSpacing")):
        list_style_ppr.remove(node)

    # LibreOffice can also collapse adjacent WPS list paragraphs or render a
    # list marker beside an inline screenshot.  Convert the template lists to
    # explicit, ordinary paragraphs so every item and image reserves its own
    # line consistently in Word, WPS and LibreOffice.
    reference_labels = {"参考放置图", "官服:", "steam服:", "国服:"}
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "List Bullet":
            continue
        text = paragraph.text.strip()
        contains_image = bool(paragraph._element.xpath(".//wp:inline"))
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.line_spacing = 1.15
        if contains_image:
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            continue
        if text in reference_labels:
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            continue
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-10)
        paragraph.paragraph_format.space_after = Pt(
            6 if text.startswith("Bundles2 模式会更新") else 3
        )
        for run in paragraph.runs:
            if run.text:
                run.text = "• " + run.text
                break

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip().startswith("POE2 三服合一物价补丁"):
                if paragraph.runs:
                    paragraph.runs[0].text = f"POE1 / POE2 物价补丁 {PATCH_VERSION}"
                    for run in paragraph.runs[1:]:
                        run.text = ""
                else:
                    paragraph.add_run(f"POE1 / POE2 物价补丁 {PATCH_VERSION}")

    localization_steps = [
        "选择 POE1 国际服 Steam Bundles2 或官方 GGPK；国服不会启用汉化按钮。",
        "点击“一键汉化POE1国际服”。程序每次都会从 PoEDB 推荐的 LibGGPK3 GitHub 最新 Release 下载 PoeChinese3_win-x64.exe；依次使用 ghfast.top、gh-proxy.com、gh.ddlc.top、ghproxy.it、github.boki.moe、ghproxy.net、gh.jasonzeng.dev、gh.monlor.com，最后回退 GitHub 官方源。SHA256 必须匹配，API 可用时还会严格匹配文件大小；不使用旧缓存，窗口持续显示已运行秒数和当前阶段。",
        "完成后在 POE1 选择第二个（法文）国旗；程序会尝试把 production_Config.ini 设为 language=fr。",
        "汉化会修改游戏包，请关闭游戏和启动器。游戏更新后需重新汉化并更新物价；校验或启动异常时先用 Steam 验证游戏文件。",
    ]
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if not paragraph.text.strip().startswith("POE1 国际服一键汉化"):
            continue
        paragraph._element.getparent().remove(paragraph._element)
        for following in paragraphs[index + 1 :]:
            if not re.match(r"^\d+\.\s*", following.text.strip()):
                break
            following._element.getparent().remove(following._element)
        break
    para(doc, "POE1 国际服一键汉化", 13, True, after=4)
    # Keep this short appendix on one page. The template already leaves only
    # the lower half of the final page for it, so use slightly tighter body
    # rhythm instead of orphaning the fourth step on a nearly empty page.
    nums(doc, localization_steps, size=10.5, after=2, line=1.05)

    generated = OUT_DIR / "使用文档.docx"
    doc.save(generated)
    targets = [*DOC_PATHS]
    for path in targets:
        path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(generated, path)
        print(path)
    print(generated)
    return True


def set_font(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def para(doc, text="", size=11, bold=False, color=None, align=None, after=6, line=1.2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_font(r)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_font(r)
        r.font.size = Pt(11)


def nums(doc, items, size=11, after=4, line=1.15):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = line
        r = p.add_run(f"{i}. {item}")
        set_font(r)
        r.font.size = Pt(size)


def build():
    if copy_template_doc():
        return

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = f"POE1 / POE2 物价补丁使用文档 {PATCH_VERSION}"
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"POE1 / POE2 物价补丁使用文档 {PATCH_VERSION}")
    set_font(r)
    r.font.size = Pt(20)
    r.font.bold = True

    for line in [
        "重要提示：打补丁会修改游戏文件，存在封号或校验风险。",
        "请在关闭游戏后运行，并自行确认可以接受风险。",
        "更新和还原都会自动识别 POE1 / POE2、国服 WeGame、国际服官方 GGPK、国际服 Steam/Epic Bundles2，也可以手动选择游戏目录。",
    ]:
        para(doc, line, 14, True, (192, 0, 0), WD_ALIGN_PARAGRAPH.CENTER, 2, 1.0)

    para(doc, "一、发布版怎么放", 13, True, after=4)
    para(doc, "“物价补丁”文件夹可以放在任意位置。放在游戏根目录下仍可最快识别；放在其它位置时，程序会检查最近目录和各平台游戏库，也可以手动浏览。")
    bullets(
        doc,
        [
            r"<POE2游戏根目录>\Content.ggpk",
            r"<POE2游戏根目录>\Bundles2\_.index.bin",
            r"<POE2游戏根目录>\物价补丁\物价补丁.exe",
        ],
    )

    para(doc, "二、选择与自动识别游戏目录", 13, True, after=4)
    bullets(
        doc,
        [
            "自动模式依次检查补丁文件夹上一级、环境变量、最近有效目录、GGG 官服注册表、已安装程序、WeGame/Steam 游戏库、Epic 清单和常见位置。",
            "国服会匹配“流放之路：降临”和 WeGameApps\\rail_apps 游戏库；目录还必须包含 Bundles2\\_.index.bin，并通过 WeGame/Rail/Tencent 特征确认。",
            "检测到 Content.ggpk 时按国际服官方 GGPK 处理；检测到非国服特征的 Bundles2 时按国际服 Steam/Epic 处理。",
            "如果发现多个客户端，程序不会猜测，请切换到手动选择。确认后的有效目录会保存，更新与还原下次可复用。POE1 还可选择自动识别、汉化补丁、简体中文、繁体中文或跟随游戏配置。",
        ],
    )

    para(doc, "三、统一操作入口", 13, True, after=4)
    nums(
        doc,
        [
            "关闭游戏。",
            "双击“物价补丁.exe”。",
            "在统一窗口中确认 POE1 / POE2、客户端路径和客户端类型；POE1 还需确认显示语言，更新时还需选择补丁范围。",
            "更新时选择价格赛季：程序从 poe2scout 的 POE1/POE2 realm 目录实时读取全部软核赛季，默认选中标记为 IsCurrent 的最新赛季；点击“刷新”可发现新赛季。历史赛季使用独立的赛季标识和缓存，POE1/POE2 国服只支持当前赛季。",
            "点击底部“开始/更新物价补丁”执行更新，或点击“还原物价补丁”恢复原版；更新会提取英文表和当前客户端目标语言表，还原会直接使用已验证的还原包或逻辑还原基线。",
            "POE1 安装第三方汉化补丁时保留“自动识别”或选择“汉化补丁”；自动模式会检查最新客户端日志，检测到中文区域名后写入繁中表。更新与还原必须使用同一目标语言。",
            "POE1 国际服使用 poe.ninja 主源，再由 poe2scout、PoEDB 逐物品补缺；POE1 国服使用 poecurrency.top 主源，再按 poe.ninja、poe2scout、PoEDB 的顺序补缺。国服主源已有价格不会被国际服来源覆盖。",
            "POE1 价格追加为 C/D，Chaos/Divine 比例按各来源实时计算；POE2 继续使用原有 poe2scout、poe.ninja、PoEDB 与 poecurrency.top v2 链路和 D/E 规则。",
            "国服优先使用 latest_buy1 / latest_sell1 最新盘口价，缺失时回退到 buy_avg / sell_avg；双边价差正常时取几何均值，差距过大时取较低一侧以降低过期均价和 OCR 异常价影响。",
            "同一个显示名对应多个游戏元数据路径时会为全部别名写入价格；国服翻译重名时使用 engname 英文别名消歧，避免漏价或串价。",
            "程序会按游戏代数、客户端类型和目标语言生成独立还原基线，工作输出再按游戏目录哈希隔离；权威副本保存在对应游戏根目录的 .poe1-price-patch 或 .poe2-price-patch。",
            "Bundles2 模式会以当前游戏包为底板刷新物价；如果已经先安装功能/词缀补丁，会尽量保留这些补丁，只清理并替换本工具写入的物价标记。",
            "连续更新只复用第一次验证通过的干净基线，不会用第二次、第三次补丁覆盖它。旧共享包只作为只读迁移输入，不再成为新写入目标。",
            "价格文件会先在隔离目录生成并校验；实时数据异常时自动尝试核心价格或当前范围和模式的兼容核心缓存。没有任何安全候选时只停止本次更新，不覆盖当前游戏。",
            "写入前会复核游戏文件未被并发修改，写入后读回目标数据核对；写入或校验失败时自动使用已验证还原包恢复。",
        ],
    )

    para(doc, "四、还原物价补丁", 13, True, after=4)
    nums(
        doc,
        [
            "关闭游戏。",
            "打开“物价补丁.exe”，确认自动识别的路径和客户端类型，或手动选择要还原的游戏目录。",
            "点击底部“还原物价补丁”执行还原；还原过程中不会抓取价格，也不会写入新的物价标记。",
            "Bundles2 模式会优先验证并使用“真实还原物价补丁.zip”恢复打补丁前的物理文件，不要求先准备 .NET。",
            "GGPK 模式会使用“还原物价补丁.zip”写回当前客户端对应的 BaseItemTypes。",
            "如果专属基线丢失或过期，程序会在临时目录只清除能够确认属于本工具的 BaseItemTypes、Words 与 EndgameMaps 标记，结构和零残留校验通过后自动重建；只有无法证明安全时才会拒绝并提示平台修复。",
            "还原写入后会读回目标数据校验；逻辑还原失败会自动重试一次，物理还原失败会恢复操作前状态。",
        ],
    )

    para(doc, "五、补丁范围", 13, True, after=4)
    bullets(
        doc,
        [
            "国服 WeGame：data/balance/simplified chinese/baseitemtypes.datc64。",
            "国际服官方 / Steam / Epic：按当前游戏语言写入，例如繁中为 data/balance/traditional chinese/baseitemtypes.datc64，英文为 data/balance/baseitemtypes.datc64。",
            "需要手动指定语言时，可设置 POE2_PATCH_LANGUAGE，例如 zh-TW、en、ja。",
            "POE1 显示语言由 GUI 保存；POE1_PATCH_LANGUAGE 仅在自动模式下继续兼容。不同目标语言使用独立缓存和还原包。",
            "Bundles2 模式会更新 _.index.bin 并写入 LibGGPK3 增量包；_.index.bin 更新时间变化是正常现象。",
            "Bundles2 模式不会直接覆盖 Tiny*.bundle.bin。",
            "Bundles2 还原优先恢复安装物价补丁前备份的 _.index.bin 和 LibGGPK3 状态；旧版本迁移生成的基线会恢复到清除本工具价格/岛屿标记且保留其它兼容补丁的状态。",
            "如果其他补丁没有改同一个 BaseItemTypes 资源，通常不会被本工具影响。",
            "如果其他补丁也改了同一个 BaseItemTypes 资源，最后写入者会覆盖同资源内对应字段。",
        ],
    )

    para(doc, "六、文件说明", 13, True, after=4)
    bullets(
        doc,
        [
            "物价补丁.exe：统一 GUI 入口；可选择更新或还原、POE1 或 POE2、自动识别或手动路径。",
            "物价补丁.zip：运行时生成的当前物价补丁包。",
            "POE1/POE2 专属还原包：文件名包含客户端类型和目标语言，工作目录还按游戏路径隔离。",
            "Bundles2 真实还原包：恢复写入前的物理索引和 LibGGPK3；持久副本位于对应游戏根目录的隐藏补丁目录。",
            r"tools\dotnet-runtime：内置 .NET 8.0.28 runtime，不要删除；发布包另含校验过的离线修复包，正常使用不需要联网，只有两者都不可用时才访问 Microsoft 源。",
            r"tools\python：内置 Python 和依赖，不要删除；自动修复时依次尝试华为云、阿里云、南京大学镜像和 Python 官方备用源。",
            "一键安装特殊补丁工具：底层写入工具目录，不要删除。",
        ],
    )

    para(doc, "七、常见问题", 13, True, after=4)
    bullets(
        doc,
        [
            "提示找不到游戏目录：改用手动选择，并选择直接包含 Content.ggpk 或 Bundles2\\_.index.bin 的游戏根目录；不要选择物价补丁文件夹。多个客户端必须明确选择。",
            "提取或写入失败：请先关闭游戏和可能占用文件的工具。",
            "提示专属基线缺失或过期：新版会先核验旧包，再离线清除本工具标记并重建；只有 DAT 损坏、结构无法解析、存在未知改动或平台并发写入时才要求验证/修复，拒绝前不会修改真实游戏。",
            "缺少价格：可能是当前价格源暂无该物品数据，或物品名无法匹配本地物品表。POE1 传奇写成 传奇名[<<价格>>]，POE2 写成 [价格|传奇名]；易刷与 PoE Overlay II 会按对应格式清理价格后使用精确传奇名查询。",
            "POE1 使用 Steam 汉化补丁但没有价格：把“POE1 显示语言”设为自动识别或汉化补丁，然后重新更新。",
            "价格源暂时不可用：程序会优先继续使用兼容缓存；如果没有安全缓存，会保留当前补丁并停止本次更新，可稍后重试。",
            "杀软报毒：自制 exe、加密脚本和修改游戏文件都可能触发敏感提示，需要自行判断风险。",
        ],
    )

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"POE1 / POE2 物价补丁 {PATCH_VERSION}")
    set_font(fr)
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    temp = OUT_DIR / "使用文档.docx"
    if temp.exists():
        temp.unlink()
    doc.save(temp)
    for path in DOC_PATHS:
        path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(temp, path)
        print(path)


if __name__ == "__main__":
    build()
